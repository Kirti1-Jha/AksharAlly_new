"""
DOCX Processor — extracts text from Microsoft Word (.docx) documents.

Uses python-docx to read all paragraphs and table cells in document order.
"""

import io
from docx import Document

from .response import build_response, build_error

MAX_SIZE_MB = 10


def process_docx(file_storage) -> dict:
    """
    Accept a Werkzeug FileStorage object for a .docx file.
    Reads all paragraphs and table cells, joins them, and returns a
    unified response.
    """
    if file_storage is None or file_storage.filename == "":
        return build_error("docx", "No DOCX file provided.")

    filename = file_storage.filename.lower()
    if not filename.endswith(".docx"):
        return build_error("docx", "File does not have a .docx extension.")

    docx_bytes = file_storage.read()

    if not docx_bytes:
        return build_error("docx", "Uploaded DOCX file is empty.")

    size_mb = len(docx_bytes) / (1024 * 1024)
    if size_mb > MAX_SIZE_MB:
        return build_error("docx", f"File too large ({size_mb:.1f} MB). Max is {MAX_SIZE_MB} MB.")

    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as e:
        return build_error("docx", f"Could not open DOCX file: {str(e)}")

    text_parts = []

    # Extract from paragraphs (main body text)
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text_parts.append(text)

    # Extract from tables (each cell is treated as a paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in text_parts:
                    text_parts.append(text)

    combined = "\n".join(text_parts).strip()

    if not combined:
        return build_error("docx", "No text could be extracted from the DOCX file.")

    return build_response("docx", combined)
